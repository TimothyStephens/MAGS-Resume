import docx
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import markdown
from bs4 import BeautifulSoup, NavigableString

def extract_text_from_docx(file_path: str) -> str:
    """Reads all text from a .docx file."""
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def add_hyperlink(paragraph, url: str, text: str, bold: bool, italic: bool):
    """
    Adds a hyperlink to a paragraph, preserving formatting.
    """
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    run = paragraph.add_run()
    run._r.append(hyperlink)
    
    # Try to apply the 'Hyperlink' style which is standard in Word
    try:
        run.style = 'Hyperlink'
    except KeyError:
        # Fallback to manual formatting if 'Hyperlink' style is not found
        run.font.color.rgb = RGBColor(0, 0, 255)
        run.font.underline = True

    run.bold = bold
    run.italic = italic
    return run

def _process_soup_element(paragraph, element, bold=False, italic=False):
    """Recursively adds runs to a paragraph based on soup tags (b, i, strong, em, a)."""
    # Inherit and update formatting from the current element's tag.
    if not isinstance(element, NavigableString):
        bold = bold or element.name in ['strong', 'b']
        italic = italic or element.name in ['em', 'i']

    if isinstance(element, NavigableString):
        text = str(element)
        
        # Normalize whitespace: convert newlines to spaces. 
        # We rely on 'nl2br' extension and <br> tag handling for explicit breaks.
        text = text.replace('\n', ' ').replace('\r', '')

        # Strip leading whitespace if this is the start of the paragraph
        if len(paragraph.runs) == 0:
            text = text.lstrip()

        # Strip leading whitespace if the previous element was a hard break <br>
        # This fixes the issue where 'nl2br' adds \n after <br>, creating a visual space at start of line
        if element.previous_sibling and getattr(element.previous_sibling, 'name', None) == 'br':
            text = text.lstrip()

        if text:
            run = paragraph.add_run(text)
            run.bold = bold
            run.italic = italic
        return

    # Handle hyperlinks before iterating children
    if element.name == 'a' and element.get('href'):
        url = element.get('href')
        # Get text, but don't recurse further into this branch to avoid nested formatting in links
        text = element.get_text()
        if text:
            add_hyperlink(paragraph, url, text, bold, italic)
        return

    # Handle explicit line breaks (from nl2br or <br> tags)
    if element.name == 'br':
        # If inside a list, treat break as a space (wrapping), otherwise real break
        is_list = paragraph.style and hasattr(paragraph.style, 'name') and 'List' in paragraph.style.name
        if is_list:
            paragraph.add_run(' ')
        else:
            paragraph.add_run().add_break()
        return

    # Skip block-level elements that are handled by the main loop or recursive list processor
    # to prevent flattening nested content into the parent paragraph.
    if element.name in ['ul', 'ol']:
        return

    for child in element.children:
        # Pass the updated bold/italic state down to children.
        _process_soup_element(paragraph, child, bold=bold, italic=italic)

def save_text_to_docx(text: str, output_path: str):
    """Saves markdown/HTML mixed text to a formatted .docx file."""
    # Pre-process: Fix common Markdown spacing issues (lists/HRs missing newlines)
    lines = text.split('\n')
    fixed_lines = []
    for i, line in enumerate(lines):
        stripped = line.strip()
        
        # Detect list item markers (including raw bullets often used in copy-paste)
        is_list_item = stripped.startswith(('* ', '- ', '•'))
        
        if i > 0:
            prev_stripped = lines[i-1].strip()
            prev_is_list_item = prev_stripped.startswith(('* ', '- ', '•'))
            
            # 1. Ensure blank line before a list starts (Text -> List)
            if is_list_item and prev_stripped and not prev_is_list_item:
                fixed_lines.append('')
            
            # 2. Ensure blank line after a list ends (List -> Text)
            elif not is_list_item and stripped and prev_is_list_item:
                fixed_lines.append('')
                
            # 3. Ensure blank line before HR markers
            elif stripped in ('---', '***', '___') and prev_stripped:
                fixed_lines.append('')
            
        # 4. Process the line content
        if not stripped:
            fixed_lines.append('\n<p>&nbsp;</p>\n')
        elif stripped.startswith('•'):
            # Normalize raw bullet to markdown asterisk to ensure correct list parsing
            # Replace the first occurrence of bullet with asterisk
            fixed_lines.append(line.replace('•', '*', 1))
        else:
            fixed_lines.append(line)
            
    text = '\n'.join(fixed_lines)

    # 1. Convert Markdown (and pass-through HTML) to full HTML structure
    html_content = markdown.markdown(text, extensions=['extra', 'nl2br'])
    
    # 2. Parse HTML structure
    soup = BeautifulSoup(html_content, 'html.parser')
    
    doc = docx.Document()
    
    # Set page margins to "Narrow" (0.5 inches)
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # --- Style Compaction ---
    # Modify the default 'Normal' style for body text
    style = doc.styles['Normal']
    p_fmt = style.paragraph_format
    p_fmt.space_before = Pt(0)
    p_fmt.space_after = Pt(4) # Slight space to distinguish paragraphs (strict spacing adherence)
    p_fmt.line_spacing = 1.0  # Single line spacing

    # Modify list styles for compactness
    for style_name in ['List Bullet', 'List Number']:
        try:
            list_style = doc.styles[style_name]
            p_fmt_list = list_style.paragraph_format
            p_fmt_list.space_before = Pt(0)
            p_fmt_list.space_after = Pt(0)
            p_fmt_list.line_spacing = 1.0
        except KeyError:
            pass # Style might not exist in the default template

    # Modify heading styles to reduce space around them
    for i in range(1, 7):
        try:
            heading_style = doc.styles[f'Heading {i}']
            p_fmt_heading = heading_style.paragraph_format
            p_fmt_heading.space_before = Pt(6) # A bit of space before
            p_fmt_heading.space_after = Pt(2)  # Very little space after
        except KeyError:
            continue # Style might not exist

    def process_list_tag(doc, list_tag, level=0):
        """Recursively handles nested lists."""
        style_name = 'List Number' if list_tag.name == 'ol' else 'List Bullet'
        
        for li in list_tag.find_all('li', recursive=False):
            p = doc.add_paragraph(style=style_name)
            
            # Set indentation for all list levels.
            # level 0 = top level, level 1 = first sublist, etc.
            # Push bullet to 0.25" (Text at 0.5", Hanging at -0.25") to avoid margin alignment
            p.paragraph_format.left_indent = Inches(0.25 * (level + 2))
            p.paragraph_format.first_line_indent = Inches(-0.25)
            
            for child in li.children:
                if child.name in ['ul', 'ol']:
                    process_list_tag(doc, child, level + 1)
                else:
                    _process_soup_element(p, child)

    # 3. Traverse HTML top-level elements and create Docx elements
    for tag in soup.find_all(recursive=False):
        # Headings (h1 - h6)
        if tag.name and tag.name.startswith('h') and len(tag.name) == 2 and tag.name[1].isdigit():
            level = int(tag.name[1])
            # add_heading returns a paragraph object
            heading = doc.add_heading(level=level)
            _process_soup_element(heading, tag)

            # Force headers to be black instead of default Word blue
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
            
            if tag.get('align') == 'center':
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Paragraphs <p>
        elif tag.name == 'p':
            # Check if this paragraph looks like a list item that Markdown missed
            # (e.g. "* Item" or "- Item")
            text_preview = tag.get_text().strip()
            
            p = doc.add_paragraph()
            _process_soup_element(p, tag)
            if tag.get('align') == 'center':
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Divs (special handling for flex/space-between headers often used in resumes)
        elif tag.name == 'div':
            style = tag.get('style', '').lower()
            # Get direct element children (ignoring whitespace nav strings)
            children = list(tag.find_all(recursive=False))
            
            p = doc.add_paragraph()
            
            # Heuristic: If flex and exactly 2 children, assume Left/Right split
            if 'flex' in style and len(children) == 2:
                # Add a right-aligned tab stop at the right margin (7.5" for Letter with 0.5" margins)
                tab_stops = p.paragraph_format.tab_stops
                tab_stops.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
                
                _process_soup_element(p, children[0]) # Left content
                p.add_run('\t')                       # Tab to right
                _process_soup_element(p, children[1]) # Right content
            else:
                # Standard div processing
                _process_soup_element(p, tag)

        # Unordered Lists <ul>
        elif tag.name == 'ul':
            process_list_tag(doc, tag)
        
        # Ordered Lists <ol>
        elif tag.name == 'ol':
            process_list_tag(doc, tag)

        # Horizontal Rules <hr>
        elif tag.name == 'hr':
            # Add a real bottom border to create a horizontal line
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'auto')
            pBdr.append(bottom)
            pPr.append(pBdr)

        else:
            # Fallback for text nodes or unknown blocks
            if tag.name:
                p = doc.add_paragraph()
                _process_soup_element(p, tag)
    
    doc.save(output_path)